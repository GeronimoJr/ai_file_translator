import streamlit as st
import requests
import tempfile
import os
import re
import traceback
import json
import time
from datetime import datetime
import pandas as pd
import io
import xml.etree.ElementTree as ET
import tiktoken
import concurrent.futures

# Warunkowy import dla Google Drive
try:
    from pydrive2.auth import GoogleAuth
    from pydrive2.drive import GoogleDrive
    from oauth2client.service_account import ServiceAccountCredentials
    GOOGLE_DRIVE_AVAILABLE = True
except ImportError:
    GOOGLE_DRIVE_AVAILABLE = False

# Warunkowy import dla wykrywania języka
try:
    import langid
    LANGID_AVAILABLE = True
except ImportError:
    LANGID_AVAILABLE = False

# Warunkowy import dla dokumentów Word
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Stałe konfiguracyjne
SUPPORTED_FILE_TYPES = ["xml", "csv", "xls", "xlsx"]
if DOCX_AVAILABLE:
    SUPPORTED_FILE_TYPES.extend(["doc", "docx"])

SUPPORTED_LANGUAGES = {
    "auto": "Automatyczne wykrywanie",
    "en": "angielski", 
    "pl": "polski", 
    "de": "niemiecki", 
    "fr": "francuski", 
    "es": "hiszpański", 
    "it": "włoski"
}
CHUNK_TOKEN_LIMIT = 10000

MODEL_PRICES = {
    "openai/gpt-4o-mini": {"prompt": 0.15, "completion": 0.6},
    "openai/gpt-4o": {"prompt": 2.5, "completion": 10.0},
    "anthropic/claude-3.5-haiku": {"prompt": 0.8, "completion": 4.0},
    "anthropic/claude-3.7-sonnet": {"prompt": 3.0, "completion": 15.0},
    "google/gemini-2.5-pro-preview": {"prompt": 1.25, "completion": 10.0}
}

# Inicjalizacja stanu sesji
def init_session_state():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "output_bytes" not in st.session_state:
        st.session_state.output_bytes = None
    if "translation_in_progress" not in st.session_state:
        st.session_state.translation_in_progress = False
    if "translation_done" not in st.session_state:
        st.session_state.translation_done = False
    if "raw_bytes" not in st.session_state:
        st.session_state.raw_bytes = None
    if "file_type" not in st.session_state:
        st.session_state.file_type = None
    if "detected_lang" not in st.session_state:
        st.session_state.detected_lang = None
    if "translated_df" not in st.session_state:
        st.session_state.translated_df = None
    if "original_df" not in st.session_state:
        st.session_state.original_df = None
    if "glossary" not in st.session_state:
        st.session_state.glossary = {}

@st.cache_data(ttl=3600)
def clean_invalid_xml_chars(text):
    return ''.join(
        c for c in text
        if c in ('\t', '\n', '\r') or
        (0x20 <= ord(c) <= 0xD7FF) or
        (0xE000 <= ord(c) <= 0xFFFD) or
        (0x10000 <= ord(c) <= 0x10FFFF)
    )

# Usunięto dekorator cache_data, ponieważ ElementTree.Element nie jest hashowalny
def parse_xml_with_fallback(raw_bytes):
    match = re.search(br'<\?xml[^>]*encoding=["\']([^"\']+)["\']', raw_bytes)
    declared_enc = match.group(1).decode('ascii').lower() if match else None
    candidates = [declared_enc] if declared_enc else []
    candidates += ['utf-16', 'utf-8', 'iso-8859-2', 'windows-1250']

    for enc in candidates:
        try:
            decoded = raw_bytes.decode(enc)
            cleaned = clean_invalid_xml_chars(decoded)
            tree = ET.ElementTree(ET.fromstring(cleaned))
            return tree, tree.getroot(), enc
        except Exception:
            continue

    return None, None, None

# Usunięto dekorator cache_data, ponieważ ElementTree.Element nie jest hashowalny
def extract_xml_texts_and_paths(elem, path=""):
    texts = []
    if elem.text and elem.text.strip():
        texts.append((f"{path}/text", elem.text.strip()))
    for k, v in elem.attrib.items():
        texts.append((f"{path}/@{k}", v))
    for i, child in enumerate(elem):
        child_path = f"{path}/{child.tag}[{i}]"
        texts.extend(extract_xml_texts_and_paths(child, child_path))
    return texts

def insert_translations_into_xml(elem, translations, path=""):
    if elem.text and elem.text.strip():
        key = f"{path}/text"
        if key in translations:
            elem.text = translations[key]
    for k in elem.attrib:
        key = f"{path}/@{k}"
        if key in translations:
            elem.attrib[k] = translations[key]
    for i, child in enumerate(elem):
        child_path = f"{path}/{child.tag}[{i}]"
        insert_translations_into_xml(child, translations, child_path)

@st.cache_data(ttl=3600)
def is_numeric_value(value):
    """Sprawdza czy wartość jest liczbą lub kodem produktu"""
    if isinstance(value, (int, float)):
        return True
    if isinstance(value, str):
        # Sprawdź wzorce dla kodów produktów (np. 700.KG-2)
        if re.match(r'\d{3}\.\w+-\d+', value):
            return True
        # Sprawdź czy to liczba z przecinkiem lub kropką
        if re.match(r'^\d+[,.]?\d*$', value.strip()):
            return True
    return False

@st.cache_data(ttl=3600)
def is_product_code(value):
    """Sprawdza czy wartość jest kodem produktu, który nie powinien być tłumaczony"""
    if isinstance(value, str):
        # Sprawdź wzorce dla kodów produktów (np. 700.KG-2)
        if re.match(r'\d{3}\.\w+-\d+', value) or re.match(r'\d{3}\.\w+-\d+[a-zA-Z.]+', value):
            return True
    return False

@st.cache_data(ttl=3600)
def format_number_for_locale(value, target_lang):
    """Formatuje liczby zgodnie z konwencją docelowego języka"""
    try:
        # Konwertuj do float jeśli to możliwe
        if isinstance(value, str):
            # Zamień przecinki na kropki dla konwersji w Pythonie
            value = value.replace(',', '.')
            value = float(value)
            
        # Formatuj zgodnie z lokalem
        if target_lang in ['en']:  # angielski używa kropki
            return str(value).replace(',', '.')
        else:  # inne europejskie języki używają przecinka
            return str(value).replace('.', ',')
    except (ValueError, TypeError):
        # Jeśli to nie jest liczba, zwróć oryginalną wartość
        return value

@st.cache_data(ttl=3600)
def detect_language(text):
    """Wykrywa język podanego tekstu"""
    try:
        if not text or len(text.strip()) < 5:
            return None
        if LANGID_AVAILABLE:
            lang, _ = langid.classify(text)
            return lang
        return None
    except:
        return None

@st.cache_data(ttl=3600)
def detect_source_language(texts):
    """Wykrywa główny język źródłowy na podstawie próbki tekstów"""
    if not LANGID_AVAILABLE or not texts:
        return "auto"  # Domyślnie auto-detect
        
    # Bierz próbkę 10 najdłuższych tekstów do analizy
    sample_texts = sorted([t for t in texts if isinstance(t, str) and len(t) > 10], 
                          key=len, reverse=True)[:10]
    
    if not sample_texts:
        return "auto"
    
    # Liczniki języków
    lang_counts = {}
    
    for text in sample_texts:
        detected = detect_language(text)
        if detected:
            lang_counts[detected] = lang_counts.get(detected, 0) + 1
    
    if not lang_counts:
        return "auto"
        
    # Zwróć najczęściej wykryty język
    source_lang = max(lang_counts.items(), key=lambda x: x[1])[0]
    return source_lang

@st.cache_data(ttl=3600)
def chunk_lines(lines, model_name="gpt-4", chunk_token_limit=10000):
    # Sprawdź czy tiktoken ma wsparcie dla danego modelu
    try:
        enc = tiktoken.encoding_for_model(model_name)
    except:
        # Fallback do cl100k_base jako bezpiecznej opcji
        enc = tiktoken.get_encoding("cl100k_base")
        
    chunks, current_chunk, current_tokens = [], [], 0
    for i, line in enumerate(lines):
        token_len = len(enc.encode(line))
        if current_tokens + token_len > chunk_token_limit:
            chunks.append(current_chunk)
            current_chunk, current_tokens = [], 0
        current_chunk.append((i, line))
        current_tokens += token_len
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

@st.cache_data(ttl=3600)
def estimate_cost(chunks, model_name):
    # Sprawdź czy tiktoken ma wsparcie dla danego modelu
    try:
        enc = tiktoken.encoding_for_model(model_name)
    except:
        # Fallback do cl100k_base jako bezpiecznej opcji
        enc = tiktoken.get_encoding("cl100k_base")
        
    prompt_tokens = sum(len(enc.encode(line)) for _, line in sum(chunks, []))
    completion_tokens = int(prompt_tokens * 1.2)
    pricing = MODEL_PRICES.get(model_name, {"prompt": 1.0, "completion": 1.0})
    cost_prompt = prompt_tokens / 1_000_000 * pricing["prompt"]
    cost_completion = completion_tokens / 1_000_000 * pricing["completion"]
    return prompt_tokens, completion_tokens, cost_prompt + cost_completion

def retry_api_call(func, max_retries=3, initial_backoff=1):
    retries = 0
    while retries < max_retries:
        try:
            return func()
        except Exception as e:
            retries += 1
            if retries >= max_retries:
                st.error(f"Wyczerpano limit prób ({max_retries})")
                raise
            wait_time = initial_backoff * (2 ** (retries - 1))  # Exponential backoff
            st.warning(f"Próba {retries} nieudana: {e}. Ponowienie za {wait_time}s")
            time.sleep(wait_time)

def call_translation_api(prompt, source_lang, target_lang, model, api_key):
    """Wywołuje API tłumaczenia z odpowiednimi instrukcjami"""
    if source_lang == "auto":
        system_prompt = (f"You are a precise translator. Translate the text from the source language to {target_lang}. "
                        f"Maintain the exact structure and format of the input. "
                        f"Don't change numbers, product codes, or measurements. "
                        f"Translate only text content, maintaining item numbers if present.")
    else:
        system_prompt = (f"You are a precise translator. Translate the text from {source_lang} to {target_lang}. "
                        f"Maintain the exact structure and format of the input. "
                        f"Don't change numbers, product codes, or measurements. "
                        f"Translate only text content, maintaining item numbers if present.")
    
    def make_api_call():
        res = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ]
            },
            timeout=60
        )
        res.raise_for_status()
        return res.json()
    
    result = retry_api_call(make_api_call)
    return result["choices"][0]["message"]["content"]

def translate_structured_csv_data(df, source_lang, target_lang, model, api_key, preserve_headers=True, maintain_numbers=True):
    """
    Dwuetapowe tłumaczenie danych CSV:
    1. Tłumaczenie nagłówków
    2. Tłumaczenie zawartości kolumn tekstowych
    """
    st.info("Przygotowywanie danych do tłumaczenia...")
    result_df = df.copy()
    
    # Zapisz oryginalne nagłówki
    original_headers = list(df.columns)
    
    # Krok 1: Tłumaczenie nagłówków jeśli potrzeba
    if not preserve_headers:
        headers_to_translate = []
        for header in original_headers:
            if not is_numeric_value(header) and not is_product_code(header):
                headers_to_translate.append(header)
        
        if headers_to_translate:
            st.info("Tłumaczenie nagłówków...")
            # Tworzenie specjalnego formatu dla nagłówków
            headers_prompt = "TRANSLATE THESE COLUMN HEADERS:\n"
            for i, header in enumerate(headers_to_translate):
                headers_prompt += f"{i+1}. {header}\n"
            
            headers_response = call_translation_api(headers_prompt, source_lang, target_lang, model, api_key)
            
            # Parsowanie odpowiedzi
            translated_headers = []
            for line in headers_response.split('\n'):
                line = line.strip()
                if line:
                    # Próbuj dopasować "numer. tłumaczenie"
                    match = re.match(r'^(\d+)\.\s+(.+)$', line)
                    if match:
                        idx, translated = int(match.group(1)) - 1, match.group(2)
                        if 0 <= idx < len(headers_to_translate):
                            translated_headers.append((headers_to_translate[idx], translated))
                    elif len(translated_headers) < len(headers_to_translate):
                        # Jeśli nie udało się dopasować wzorca, dodaj jako jest
                        translated_headers.append((headers_to_translate[len(translated_headers)], line))
            
            # Zastosuj tłumaczenia nagłówków
            header_map = {orig: trans for orig, trans in translated_headers}
            new_headers = [header_map.get(h, h) for h in original_headers]
            result_df.columns = new_headers

    # Krok 2: Identyfikacja kolumn tekstowych do tłumaczenia
    text_columns = []
    for col in original_headers:
        # Sprawdź, czy kolumna zawiera wartości tekstowe (niebędące liczbami/kodami)
        sample_values = df[col].dropna().astype(str).unique()[:20]
        if any(not is_numeric_value(val) and not is_product_code(val) and str(val).strip() for val in sample_values):
            text_columns.append(col)
    
    # Krok 3: Tłumaczenie zawartości kolumn tekstowych
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for col_idx, col in enumerate(text_columns):
        status_text.text(f"Tłumaczenie kolumny {col_idx+1}/{len(text_columns)}: {col}")
        
        # Pobierz wszystkie unikalne wartości tekstowe z kolumny
        unique_values = {}
        for idx, val in df[col].items():
            if pd.notna(val) and not is_numeric_value(val) and not is_product_code(val) and str(val).strip():
                val_str = str(val).strip()
                if val_str not in unique_values:
                    unique_values[val_str] = []
                unique_values[val_str].append(idx)
        
        # Jeśli są wartości do tłumaczenia
        if unique_values:
            values_list = list(unique_values.keys())
            
            # Tworzymy strukturyzowany prompt
            values_prompt = f"TRANSLATE THE FOLLOWING VALUES FROM COLUMN '{col}':\n"
            for i, val in enumerate(values_list):
                values_prompt += f"{i+1}. {val}\n"
            
            # Dodaj instrukcję o zachowaniu numeracji
            values_prompt += "\nKeep the same numbering format in your response. Report each translated item with its number."
            
            # Wywołaj API tłumaczenia
            values_response = call_translation_api(values_prompt, source_lang, target_lang, model, api_key)
            
            # Parsowanie odpowiedzi z zachowaniem numeracji
            value_translations = {}
            response_lines = values_response.split("\n")
            
            for line in response_lines:
                line = line.strip()
                if not line:
                    continue
                
                # Dopasuj wzorzec "numer. tłumaczenie"
                match = re.match(r'^(\d+)\.\s+(.+)$', line)
                if match:
                    idx, translated = int(match.group(1)) - 1, match.group(2)
                    if 0 <= idx < len(values_list):
                        original = values_list[idx]
                        value_translations[original] = translated
            
            # Jeśli parsowanie się nie powiodło, spróbuj prostego podejścia
            if not value_translations and len(response_lines) == len(values_list):
                for i, line in enumerate(response_lines):
                    if i < len(values_list):
                        value_translations[values_list[i]] = line.strip()
            
            # Zastosuj tłumaczenia do DataFrame
            col_in_result = col
            if not preserve_headers and col in header_map:
                col_in_result = header_map[col]
                
            for original, indices in unique_values.items():
                if original in value_translations:
                    for idx in indices:
                        result_df.at[idx, col_in_result] = value_translations[original]
        
        # Aktualizuj pasek postępu
        progress_bar.progress((col_idx + 1) / len(text_columns))
    
    status_text.text("Tłumaczenie zakończone!")
    return result_df

def validate_translation_results(original_df, translated_df):
    """
    Sprawdza poprawność tłumaczenia i naprawia problemy:
    1. Sprawdza strukturę
    2. Weryfikuje, że wartości liczbowe i kody nie są zmienione
    3. Sprawdza czy nie ma pustych tłumaczeń
    """
    st.info("Walidacja i naprawa tłumaczenia...")
    
    # Sprawdź strukturę
    if original_df.shape != translated_df.shape:
        st.warning(f"Niezgodna struktura: oryginał {original_df.shape}, tłumaczenie {translated_df.shape}")
        return translated_df  # Trudno naprawić różną strukturę

    # Przygotuj mapę oryginalnych do przetłumaczonych kolumn
    original_cols = list(original_df.columns)
    translated_cols = list(translated_df.columns)
    
    # Sprawdź wartości liczbowe i kody
    for i, col_orig in enumerate(original_cols):
        col_trans = translated_cols[i]
        
        for idx in original_df.index:
            orig_val = original_df.at[idx, col_orig]
            
            # Sprawdź wartości liczbowe i kody produktów
            if pd.notna(orig_val) and (is_numeric_value(orig_val) or is_product_code(orig_val)):
                trans_val = translated_df.at[idx, col_trans]
                
                # Jeśli wartość została zmieniona, przywróć oryginalną
                if str(orig_val).strip() != str(trans_val).strip():
                    translated_df.at[idx, col_trans] = orig_val
            
            # Sprawdź puste tłumaczenia
            elif pd.notna(orig_val) and str(orig_val).strip():
                trans_val = translated_df.at[idx, col_trans]
                if pd.isna(trans_val) or not str(trans_val).strip():
                    translated_df.at[idx, col_trans] = orig_val

    return translated_df

def write_csv_with_original_format(df, original_df, output_path, encoding='utf-8'):
    """Zapisuje DataFrame do CSV z zachowaniem formatu i separatorów z oryginalnego pliku"""
    # Określ separator na podstawie oryginalnego pliku
    with open(output_path, 'w', encoding=encoding) as f:
        # Zapisz nagłówki
        f.write(','.join([str(col) for col in df.columns]) + '\n')
        
        # Zapisz wiersze
        for idx in df.index:
            row_values = []
            for col_idx, col in enumerate(df.columns):
                val = df.at[idx, col]
                
                # Zachowaj oryginalny format liczb, jeśli to możliwe
                if pd.notna(val):
                    if col_idx < len(original_df.columns):
                        orig_col = original_df.columns[col_idx]
                        if idx in original_df.index:
                            orig_val = original_df.at[idx, orig_col]
                            # Jeśli mamy taką samą liczbę ale w innym formacie
                            if is_numeric_value(val) and is_numeric_value(orig_val):
                                if ',' in str(orig_val) and '.' in str(val):
                                    val = str(val).replace('.', ',')
                                elif '.' in str(orig_val) and ',' in str(val):
                                    val = str(val).replace(',', '.')
                
                # Dodaj cudzysłowy, jeśli zawiera przecinki
                if isinstance(val, str) and (',' in val or '"' in val or '\n' in val):
                    val = f'"{val.replace('"', '""')}"'
                
                row_values.append(str(val) if pd.notna(val) else '')
            
            f.write(','.join(row_values) + '\n')

@st.cache_data(ttl=3600)
def parse_csv_with_encoding_fallback(raw_bytes):
    encodings = ['utf-8', 'iso-8859-1', 'iso-8859-2', 'windows-1250']
    for enc in encodings:
        try:
            return pd.read_csv(io.BytesIO(raw_bytes), encoding=enc), enc
        except UnicodeDecodeError:
            continue
    st.error("Nie udało się rozpoznać kodowania pliku CSV")
    raise ValueError("Nieobsługiwane kodowanie pliku")

@st.cache_data(ttl=3600)
def parse_csv_with_separator_fallback(raw_bytes, encoding):
    for sep in [',', ';', '\t']:
        try:
            df = pd.read_csv(io.BytesIO(raw_bytes), encoding=encoding, sep=sep)
            if len(df.columns) > 1:  # Sprawdź czy format ma sens
                return df, sep
        except Exception:
            continue
    
    # Ostatnia próba z automatycznym wykrywaniem separatora
    try:
        df = pd.read_csv(io.BytesIO(raw_bytes), encoding=encoding, sep=None, engine='python')
        return df, ','  # Zakładamy przecinek jako domyślny separator
    except Exception as e:
        st.error(f"Nie udało się odczytać pliku CSV: {e}")
        raise

@st.cache_data(ttl=3600)
def parse_excel_file(raw_bytes):
    """Parsowanie pliku Excel z cache'owaniem"""
    return pd.read_excel(io.BytesIO(raw_bytes))

# Funkcja nie używa dekoratora cache_data, gdyż obiekty Document nie są serializable
def parse_doc_file(raw_bytes):
    """Parsowanie pliku DOC/DOCX bez cache'owania"""
    if not DOCX_AVAILABLE:
        st.error("Biblioteka python-docx nie jest dostępna. Nie można przetwarzać plików .docx")
        return None, []
    
    doc = Document(io.BytesIO(raw_bytes))
    
    # Zamiast zwracać obiekt Document, zwracamy tylko tekst
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text.strip())
    
    # Tabele
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_texts.append(cell.text.strip())
    
    # Zapamiętaj ilość paragrafów (potrzebne przy odtwarzaniu dokumentu)
    paragraph_count = len(paragraphs)
    
    # Zwróć złączony tekst
    all_texts = paragraphs + table_texts
    return {"paragraphs": paragraphs, "table_texts": table_texts}, all_texts

def translate_chunks_with_progress(chunks, source_lang, target_lang, model, api_key):
    """Wersja funkcji translate_chunks z paskiem postępu"""
    translated_pairs = []
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, chunk in enumerate(chunks):
        status_text.text(f"Tłumaczenie części {i + 1} z {len(chunks)}...")
        content = "\n".join(line for _, line in chunk)
        expected_count = len(chunk)
        
        # Utworzenie instrukcji z jasnym formatem wyjściowym
        if source_lang == "auto":
            prompt = (f"Translate the following text to {target_lang}. "
                     f"Keep exactly the same structure, preserving all numbers, codes and special characters. "
                     f"Return each line translated in the original order, without adding line numbers.\n\n"
                     f"Text to translate:\n{content}")
        else:
            prompt = (f"Translate the following text from {source_lang} to {target_lang}. "
                     f"Keep exactly the same structure, preserving all numbers, codes and special characters. "
                     f"Return each line translated in the original order, without adding line numbers.\n\n"
                     f"Text to translate:\n{content}")
        
        # Dodaj instrukcję systemową z jasnymi wytycznymi
        system_prompt = ("You are a precise translator. Translate exactly what is provided, "
                        "preserving all numbers, product codes, measurements and technical specifications. "
                        "Do not add, remove or change any numerical values. "
                        "Keep the original format intact.")
        
        def make_api_call():
            res = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={
                    "model": model, 
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": prompt}
                    ]
                },
                timeout=60
            )
            res.raise_for_status()
            return res.json()
        
        try:
            result = retry_api_call(make_api_call)
            result_lines = result["choices"][0]["message"]["content"].splitlines()
            
            # Dopasuj liczbę linii w wyniku
            if len(result_lines) < expected_count:
                st.warning(f"Brakujące linie w tłumaczeniu ({len(result_lines)} zamiast {expected_count})")
                result_lines += [""] * (expected_count - len(result_lines))
            elif len(result_lines) > expected_count:
                st.warning(f"Dodatkowe linie w tłumaczeniu ({len(result_lines)} zamiast {expected_count})")
                result_lines = result_lines[:expected_count]
            
            # Utwórz pary (indeks, tłumaczenie)
            for (idx, original), translated in zip(chunk, result_lines):
                translated_pairs.append((idx, translated.strip()))
                
        except Exception as e:
            st.error(f"Błąd podczas tłumaczenia: {e}")
            # Wstaw oryginały dla nieudanych tłumaczeń
            for idx, original in chunk:
                translated_pairs.append((idx, original))
        
        # Aktualizuj pasek postępu
        progress_bar.progress((i + 1) / len(chunks))
    
    status_text.text("Tłumaczenie zakończone!")
    
    # Sortuj według oryginalnego indeksu
    translated_pairs.sort()
    return translated_pairs

def handle_file_upload():
    """Obsługa przesłania pliku z zarządzaniem stanem"""
    uploaded_file = st.file_uploader("Wgraj plik do przetłumaczenia", type=SUPPORTED_FILE_TYPES)
    
    if uploaded_file is not None:
        # Resetuj stan jeśli przesłano nowy plik
        if "file_name" not in st.session_state or st.session_state.file_name != uploaded_file.name:
            st.session_state.file_name = uploaded_file.name
            st.session_state.file_type = uploaded_file.name.split(".")[-1].lower()
            st.session_state.raw_bytes = uploaded_file.read()
            st.session_state.translation_done = False
            st.session_state.translation_in_progress = False
            st.session_state.output_bytes = None
            st.session_state.original_df = None
            st.session_state.translated_df = None
            
        return True
    else:
        # Resetuj stan jak nie ma pliku
        if "file_name" in st.session_state:
            del st.session_state.file_name
        st.session_state.file_type = None
        st.session_state.raw_bytes = None
        st.session_state.translation_done = False
        st.session_state.translation_in_progress = False
        st.session_state.output_bytes = None
        st.session_state.original_df = None
        st.session_state.translated_df = None
        
        return False

def process_file():
    """Przetwarzanie przesłanego pliku"""
    file_type = st.session_state.file_type
    raw_bytes = st.session_state.raw_bytes
    
    try:
        if file_type == "xml":
            tree, root, encoding = parse_xml_with_fallback(raw_bytes)
            if not tree:
                st.error("Nie udało się odczytać pliku XML.")
                return None
            pairs = extract_xml_texts_and_paths(root)
            if not pairs:
                st.warning("Nie znaleziono tekstów do tłumaczenia w XML.")
                return None
            
            keys, lines = zip(*pairs) if pairs else ([], [])
            
            # Zapisz dane w stanie sesji
            st.session_state.xml_keys = keys
            st.session_state.xml_tree = tree
            st.session_state.xml_root = root
            st.session_state.xml_encoding = encoding
            
            return lines
            
        elif file_type == "csv":
            df, encoding = parse_csv_with_encoding_fallback(raw_bytes)
            df, separator = parse_csv_with_separator_fallback(raw_bytes, encoding)
            
            # Zapisz dane w stanie sesji
            st.session_state.csv_encoding = encoding
            st.session_state.csv_separator = separator
            st.session_state.original_df = df
            
            # Przygotowanie do estymacji kosztów
            texts_to_translate = []
            
            if not st.session_state.get("preserve_headers", True):
                for col in df.columns:
                    if not is_numeric_value(col) and not is_product_code(col):
                        texts_to_translate.append(str(col))
                
            for col in df.columns:
                for _, val in df[col].items():
                    if pd.notna(val) and not is_numeric_value(val) and not is_product_code(val) and str(val).strip():
                        texts_to_translate.append(str(val))
                
            return texts_to_translate
            
        elif file_type in ["xls", "xlsx"]:
            df = parse_excel_file(raw_bytes)
            
            # Zapisz dane w stanie sesji
            st.session_state.original_df = df
            
            # Przygotowanie do estymacji kosztów
            texts_to_translate = []
            
            if not st.session_state.get("preserve_headers", True):
                for col in df.columns:
                    if not is_numeric_value(col) and not is_product_code(col):
                        texts_to_translate.append(str(col))
                
            for col in df.columns:
                for _, val in df[col].items():
                    if pd.notna(val) and not is_numeric_value(val) and not is_product_code(val) and str(val).strip():
                        texts_to_translate.append(str(val))
                
            return texts_to_translate
            
        elif file_type in ["doc", "docx"]:
            if not DOCX_AVAILABLE:
                st.error("Biblioteka python-docx nie jest dostępna. Nie można przetwarzać plików .docx")
                return None
                
            doc_data, lines = parse_doc_file(raw_bytes)
            
            # Zapisz dane w stanie sesji
            st.session_state.doc_data = doc_data
                        
            return lines
        else:
            st.error("Nieobsługiwany typ pliku.")
            return None
    
    except Exception as e:
        st.error(f"Błąd podczas przetwarzania pliku: {e}")
        st.exception(e)
        return None

def save_translation_to_file(output_path, file_type):
    """Zapisuje przetłumaczony plik na dysk"""
    if file_type == "xml":
        tree = st.session_state.xml_tree
        encoding = st.session_state.xml_encoding or "utf-8"
        tree.write(output_path, encoding=encoding, xml_declaration=True)
    
    elif file_type in ["csv"]:
        translated_df = st.session_state.translated_df
        original_df = st.session_state.original_df
        encoding = st.session_state.csv_encoding
        write_csv_with_original_format(translated_df, original_df, output_path, encoding)
        
    elif file_type in ["xls", "xlsx"]:
        translated_df = st.session_state.translated_df
        translated_df.to_excel(output_path, index=False)
        
    elif file_type in ["doc", "docx"]:
        if not DOCX_AVAILABLE:
            st.error("Biblioteka python-docx nie jest dostępna. Nie można zapisać pliku .docx")
            return
            
        # Tworzymy nowy dokument z przetłumaczonym tekstem
        new_doc = Document()
        
        # Dodajemy paragrafy
        doc_data = st.session_state.doc_data
        translated_texts = st.session_state.translated_texts
        
        # Dodajemy paragrafy
        para_count = len(doc_data["paragraphs"])
        for i in range(para_count):
            new_doc.add_paragraph(translated_texts[i])
            
        # Dodajemy tekst tabel
        table_count = len(doc_data["table_texts"])
        if table_count > 0:
            # Dodajemy separator
            new_doc.add_paragraph("---")
            
            # Dodajemy przetłumaczony tekst tabel
            for i in range(table_count):
                idx = para_count + i
                if idx < len(translated_texts):
                    new_doc.add_paragraph(translated_texts[idx])
        
        # Zapisz dokument
        new_doc.save(output_path)

def save_to_google_drive(output_path, file_type):
    """Zapisuje plik na Google Drive"""
    if not GOOGLE_DRIVE_AVAILABLE:
        st.warning("Biblioteka pydrive2 nie jest dostępna. Zapis na Google Drive niemożliwy.")
        return False
    
    drive_folder_id = st.secrets.get("GOOGLE_DRIVE_FOLDER_ID")
    service_account_json = st.secrets.get("GOOGLE_DRIVE_CREDENTIALS_JSON")
    
    if not drive_folder_id or not service_account_json:
        st.warning("Brak konfiguracji dla Google Drive. Sprawdź secrets.toml")
        return False
    
    try:
        creds_dict = json.loads(service_account_json)
        scope = ["https://www.googleapis.com/auth/drive"]
        credentials = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        gauth = GoogleAuth()
        gauth.credentials = credentials
        drive = GoogleDrive(gauth)
        now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        result_filename = f"translated_output.{now}.{file_type}"
        result_file = drive.CreateFile({"title": result_filename, "parents": [{"id": drive_folder_id}]})
        result_file.SetContentFile(output_path)
        result_file.Upload()
        st.success("Plik zapisany na Twoim Google Drive ✅")
        return True
    except Exception as e:
        st.error(f"Błąd podczas zapisu na Google Drive: {e}")
        return False

def start_translation():
    """Rozpocznij proces tłumaczenia"""
    st.session_state.translation_in_progress = True

def handle_translation():
    """Obsługa procesu tłumaczenia"""
    file_type = st.session_state.file_type
    
    try:
        if file_type in ["csv", "xls", "xlsx"]:
            # Tłumaczenie dla plików tabelarycznych
            df = st.session_state.original_df
            source_lang = st.session_state.source_lang
            if source_lang == "auto" and st.session_state.detected_lang:
                source_lang = st.session_state.detected_lang
                
            target_lang = st.session_state.target_lang
            model = st.session_state.model
            api_key = st.secrets["OPENROUTER_API_KEY"]
            preserve_headers = st.session_state.get("preserve_headers", True)
            maintain_numbers = st.session_state.get("maintain_numbers", True)
            
            # Tłumaczenie tabeli
            translated_df = translate_structured_csv_data(
                df, source_lang, target_lang, model, api_key, 
                preserve_headers=preserve_headers,
                maintain_numbers=maintain_numbers
            )
            
            # Walidacja i naprawa problemów
            translated_df = validate_translation_results(df, translated_df)
            
            # Zapisz wynik w stanie sesji
            st.session_state.translated_df = translated_df
            
            with tempfile.TemporaryDirectory() as tmpdir:
                output_path = os.path.join(tmpdir, f"output.{file_type}")
                save_translation_to_file(output_path, file_type)
                    
                with open(output_path, "rb") as f:
                    st.session_state.output_bytes = f.read()
                
                # Opcjonalnie zapisz na Google Drive
                save_to_google_drive(output_path, file_type)
            
        else:
            # Tłumaczenie dla XML i dokumentów
            chunks = st.session_state.chunks
            source_lang = st.session_state.source_lang
            if source_lang == "auto" and st.session_state.detected_lang:
                source_lang = st.session_state.detected_lang
            target_lang = st.session_state.target_lang
            model = st.session_state.model
            api_key = st.secrets["OPENROUTER_API_KEY"]
            
            translated_pairs = translate_chunks_with_progress(chunks, source_lang, target_lang, model, api_key)
            
            with tempfile.TemporaryDirectory() as tmpdir:
                output_path = os.path.join(tmpdir, f"output.{file_type}")
                
                if file_type == "xml":
                    keys = st.session_state.xml_keys
                    root = st.session_state.xml_root
                    translated_map = {keys[i]: line for i, (_, line) in enumerate(translated_pairs)}
                    insert_translations_into_xml(root, translated_map)
                    st.session_state.xml_tree.write(output_path, encoding="utf-8", xml_declaration=True)
                
                elif file_type in ["doc", "docx"]:
                    if not DOCX_AVAILABLE:
                        st.error("Biblioteka python-docx nie jest dostępna. Nie można przetworzyć pliku .docx")
                        st.session_state.translation_in_progress = False
                        return
                        
                    # Zapisz przetłumaczone teksty w porządku oryginalnym
                    translated_texts = []
                    for idx, text in translated_pairs:
                        translated_texts.append(text)
                        
                    # Zapisz przetłumaczone teksty do użycia przy zapisie
                    st.session_state.translated_texts = translated_texts
                    
                    # Zapisz przetłumaczony dokument
                    save_translation_to_file(output_path, file_type)
                
                with open(output_path, "rb") as f:
                    st.session_state.output_bytes = f.read()
                
                # Opcjonalnie zapisz na Google Drive
                save_to_google_drive(output_path, file_type)
        
        st.session_state.translation_done = True
        st.session_state.translation_in_progress = False
            
    except Exception as e:
        st.error(f"Błąd podczas tłumaczenia: {e}")
        st.exception(e)
        st.session_state.translation_in_progress = False

def run_streamlit_app():
    # Inicjalizacja stanu sesji
    init_session_state()
    
    st.set_page_config(page_title="Tłumacz plików AI", layout="centered")
    st.title("AI Tłumacz plików CSV, XML, Excel i Word")
    
    st.markdown("""
    To narzędzie umożliwa tłumaczenie zawartości plików CSV, XML, XLS, XLSX, DOC i DOCX za pomocą wybranego modelu LLM.
    Prześlij plik, wybierz język źródłowy i docelowy oraz model.
    """)
    
    # Uwierzytelnianie
    if not st.session_state.authenticated:
        user = st.text_input("Login")
        password = st.text_input("Hasło", type="password")
        if st.button("Zaloguj"):
            try:
                if user == st.secrets.get("APP_USER") and password == st.secrets.get("APP_PASSWORD"):
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("Nieprawidłowy login lub hasło")
            except Exception as e:
                st.error(f"Błąd uwierzytelniania: {e}")
                st.error("Sprawdź konfigurację secrets.toml")
        return
    
    # Interfejs główny pobierania pliku
    file_uploaded = handle_file_upload()
    
    if not file_uploaded:
        return
    
    # Przetwarzanie pliku i wyświetlenie opcji
    lines = process_file()
    
    if lines is None:
        return
    
    # Opcje dla CSV/Excel
    if st.session_state.file_type in ["csv", "xls", "xlsx"]:
        st.session_state.preserve_headers = st.checkbox("Zachowaj oryginalne nagłówki", value=True)
        st.session_state.maintain_numbers = st.checkbox("Zachowaj oryginalne wartości liczbowe", value=True)
    
    # Wybór języka źródłowego i docelowego
    st.session_state.source_lang = st.selectbox(
        "Język źródłowy", 
        list(SUPPORTED_LANGUAGES.keys()), 
        format_func=lambda x: f"{x} - {SUPPORTED_LANGUAGES[x]}" if x != "auto" else SUPPORTED_LANGUAGES[x],
        index=0  # Domyślnie "auto"
    )
    
    st.session_state.target_lang = st.selectbox(
        "Język docelowy", 
        [lang for lang in SUPPORTED_LANGUAGES.keys() if lang != "auto"], 
        format_func=lambda x: f"{x} - {SUPPORTED_LANGUAGES[x]}"
    )
    
    st.session_state.model = st.selectbox(
        "Wybierz model LLM (OpenRouter)", 
        list(MODEL_PRICES.keys()) + ["openai/gpt-4o-mini", "openai/gpt-4o", "anthropic/claude-3.5-haiku", "anthropic/claude-3.7-sonnet", "google/gemini-2.5-pro-preview"]
    )
    
    # Wykryj język, jeśli ustawiony na auto
    if st.session_state.source_lang == "auto" and lines and LANGID_AVAILABLE:
        detected_lang = detect_source_language(lines)
        st.session_state.detected_lang = detected_lang
        st.info(f"Wykryto język źródłowy: {detected_lang}")
    
    # Przygotowanie chunków i estymacja kosztów
    st.session_state.chunks = chunk_lines(lines, model_name="gpt-4", chunk_token_limit=CHUNK_TOKEN_LIMIT)
    chunks = st.session_state.chunks
    prompt_tokens, completion_tokens, cost_total = estimate_cost(chunks, st.session_state.model)
    
    st.info(f"Szacunkowe zużycie tokenów: ~{prompt_tokens} (prompt) + ~{completion_tokens} (output)")
    st.info(f"Szacunkowy koszt tłumaczenia: ~${cost_total:.4f} USD")
    
    # Obsługa tłumaczenia
    if not st.session_state.translation_in_progress and not st.session_state.translation_done:
        if st.button("Przetłumacz plik"):
            start_translation()
            st.rerun()
    
    # Tłumaczenie w trakcie
    if st.session_state.translation_in_progress:
        handle_translation()
        # Po zakończeniu tłumaczenia, odśwież interfejs
        if st.session_state.translation_done:
            st.rerun()
    
    # Wynik tłumaczenia
    if st.session_state.translation_done:
        st.success("Tłumaczenie zakończone. Plik gotowy do pobrania.")
        
        # Wyświetl przykładowe dane dla plików tabelarycznych
        if st.session_state.file_type in ["csv", "xls", "xlsx"]:
            # Porównanie oryginału i tłumaczenia
            col1, col2 = st.columns(2)
            with col1:
                st.write("Przykładowe dane oryginalne:")
                st.dataframe(st.session_state.original_df.head(5))
            with col2:
                st.write("Przykładowe dane przetłumaczone:")
                st.dataframe(st.session_state.translated_df.head(5))
        
        # Przycisk do pobrania
        if st.session_state.output_bytes:
            st.download_button(
                "📁 Pobierz przetłumaczony plik", 
                data=st.session_state.output_bytes, 
                file_name=f"translated_output.{st.session_state.file_type}", 
                mime="application/octet-stream"
            )
        
        # Status zapisania na Google Drive
        if GOOGLE_DRIVE_AVAILABLE:
            st.info("Plik został automatycznie zapisany na Google Drive (jeśli skonfigurowano)")
        
        # Opcja do resetowania i rozpoczęcia nowego tłumaczenia
        if st.button("Rozpocznij nowe tłumaczenie"):
            st.session_state.translation_done = False
            st.session_state.translation_in_progress = False
            st.session_state.output_bytes = None
            st.rerun()

if __name__ == "__main__":
    run_streamlit_app()
